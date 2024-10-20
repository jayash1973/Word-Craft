import streamlit as st
import requests
import json
import os
import pyaudio as pyaudio
import base64
from collections import defaultdict
from io import BytesIO
import speech_recognition as sr
import pyttsx3
from together import Together
import tempfile
from audio_recorder_streamlit import audio_recorder
from PIL import Image
import random
import string
import av
from bs4 import BeautifulSoup
import time
import io
import queue
from docx import Document
from fpdf import FPDF
import pytesseract
import plotly.express as px
import pandas as pd
from pdf2image import convert_from_path
import PyPDF2
import docx2txt
from streamlit_option_menu import option_menu
import nltk
from nltk.tokenize import sent_tokenize
import networkx as nx
import speech_recognition as sr
from streamlit_webrtc import webrtc_streamer, WebRtcMode, RTCConfiguration
import matplotlib.pyplot as plt
import plotly.graph_objects as go
from plotly.subplots import make_subplots

nltk.download('punkt')

from streamlit_agraph import agraph, Node, Edge, Config
import pandas as pd
from wordcloud import WordCloud
import numpy as np
import threading
import pydub
from datetime import datetime
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
import math
import wave
import re
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph
from reportlab.lib.styles import getSampleStyleSheet
import re
from dotenv import load_dotenv

load_dotenv()

TOGETHER_API_KEY = os.getenv("TOGETHER_API_KEY")

together_client = Together(api_key=TOGETHER_API_KEY)

engine = pyttsx3.init()

# Predefined characters (expanded)
predefined_characters = {
    "Naruto Uzumaki": {
        "description": "Protagonist of the Naruto series. An optimistic and determined ninja with the goal of becoming Hokage.",
        "image": "https://i.pinimg.com/originals/2e/f7/0d/2ef70d5217b530dfb766a45d9babbb5e.jpg"
        },
    "Sherlock Holmes": {
        "description": "Brilliant detective known for his logical reasoning and observational skills.",
        "image": "https://encrypted-tbn0.gstatic.com/images?q=tbn:ANd9GcSFNrsQ14wE7vKH_46oN-tCF3YwGyjo9fbLMuTpvm3ENf_10JcTHqoIyxkl_EDgpmGnEXs&usqp=CAU"
    },
    "Elizabeth Bennet": {
        "description": "Protagonist of Pride and Prejudice, known for her intelligence and wit.",
        "image": "https://www.indiependent.co.uk/wp-content/uploads/2015/08/elizabeth-bennet.jpg"
    }
}

# Initialize session state
if "characters" not in st.session_state or not st.session_state.characters:
    st.session_state.characters = predefined_characters.copy()
if "current_character" not in st.session_state:
    st.session_state.current_character = None
if "messages" not in st.session_state:
    st.session_state.messages = []
if "book_details" not in st.session_state:
    st.session_state.book_details = {}

# Fiction genres
fiction_genres = [
    "Fantasy", "Science Fiction", "Mystery", "Thriller", "Romance", "Historical Fiction",
    "Horror", "Adventure", "Contemporary Fiction", "Dystopian", "Young Adult"
]

# Non-fiction genres
non_fiction_genres = [
    "Biography", "Autobiography", "Memoir", "Self-help", "History", "Science",
    "Philosophy", "Psychology", "Business", "Travel", "True Crime"
]

def generate_ai_response(character, user_message, language):
    response = together_client.chat.completions.create(
        model="meta-llama/Meta-Llama-3.1-70B-Instruct-Turbo",
        messages=[
            {"role": "system", "content": f"You are {character}. {st.session_state.characters[character]['description']} Respond in {language}."},
            {"role": "user", "content": user_message}
        ],
        max_tokens=1024,
        temperature=0.7,
        top_p=0.7,
        top_k=50,
        repetition_penalty=1,
        stop=["<|eot_id|>", "<|eom_id|>"]
    )
    return response.choices[0].message.content

def speech_to_text(audio_bytes):
    recognizer = sr.Recognizer()
    with sr.AudioFile(BytesIO(audio_bytes)) as source:
        audio = recognizer.record(source)
    try:
        text = recognizer.recognize_google(audio)
        return text
    except sr.UnknownValueError:
        return "Sorry, I couldn't understand the audio."
    except sr.RequestError:
        return "Sorry, there was an error processing the audio."

def process_audio(frame):
    sound = frame.to_ndarray()
    sound = sound.astype(np.int16)
    return sound.tobytes()

def split_text(text, chunk_size=5000):
    """Split the text into chunks of approximately chunk_size characters."""
    chunks = []
    current_chunk = []
    current_size = 0
    for sentence in nltk.sent_tokenize(text):
        sentence_size = len(sentence)
        if current_size + sentence_size > chunk_size and current_chunk:
            chunks.append(' '.join(current_chunk))
            current_chunk = []
            current_size = 0
        current_chunk.append(sentence)
        current_size += sentence_size
    if current_chunk:
        chunks.append(' '.join(current_chunk))
    return chunks

def generate_book_chapter(book_details, chapter_index):
    chapter = book_details['chapters'][chapter_index]
    chapter_name = chapter['name']

    style_instruction = f"Use the following writing style throughout the chapter: {book_details['style_sample']}"

    if book_details['genre_type'] == 'Fiction':
        chapter_prompt = f"""Write chapter {chapter_index + 1} titled '{chapter_name}' for a {book_details['genre']} book. The book is about: {book_details['description']}. This chapter should include the following scenes: {', '.join(chapter['scenes'])}. Maintain consistency with the overall story and characters.

{style_instruction}"""
        
        chapter_content = ""
        for scene_index, scene in enumerate(chapter['scenes']):
            scene_prompt = f"""Write scene {scene_index + 1} for chapter {chapter_index + 1} titled '{chapter_name}'. The scene should include the following details: {scene}

{style_instruction}"""
            scene_text = together_client.chat.completions.create(
                model="meta-llama/Meta-Llama-3.1-70B-Instruct-Turbo",
                messages=[
                    {"role": "system", "content": f"You are an expert {book_details['genre']} writer."},
                    {"role": "user", "content": scene_prompt}
                ],
                max_tokens=6048,
                temperature=0.7
            )
            chapter_content += f"## Scene {scene_index + 1}\n\n{scene_text.choices[0].message.content}\n\n"
    else:
        chapter_prompt = f"""Write chapter {chapter_index + 1} titled '{chapter_name}' for a {book_details['genre']} non-fiction book. The book is about: {book_details['description']}. This chapter should cover the following parts: {', '.join(chapter['parts'])}. Ensure the content is informative and well-structured.

{style_instruction}"""
        
        chapter_content = ""
        for part_index, part in enumerate(chapter['parts']):
            part_prompt = f"""Write part {part_index + 1} for chapter {chapter_index + 1} titled '{chapter_name}'. The part should cover the following details: {part}

{style_instruction}"""
            part_text = together_client.chat.completions.create(
                model="meta-llama/Meta-Llama-3.1-70B-Instruct-Turbo",
                messages=[
                    {"role": "system", "content": f"You are an expert {book_details['genre']} writer."},
                    {"role": "user", "content": part_prompt}
                ],
                max_tokens=6048,
                temperature=0.7
            )
            chapter_content += f"## Part {part_index + 1}\n\n{part_text.choices[0].message.content}\n\n"
    
    return chapter_content
def generate_book_introduction(book_details):
    intro_prompt = f"""Write an engaging introduction for a {book_details['genre']} {'fiction' if book_details['genre_type'] == 'Fiction' else 'non-fiction'} book titled '{book_details['title']}' by {book_details['author']}. The book is about: {book_details['description']}.

Use the following writing style: {book_details['style_sample']}

Maintain this writing style throughout the introduction."""

    intro_content = together_client.chat.completions.create(
        model="meta-llama/Meta-Llama-3.1-70B-Instruct-Turbo",
        messages=[
            {"role": "system", "content": "You are an expert book introduction writer."},
            {"role": "user", "content": intro_prompt}
        ],
        max_tokens=1024,
        temperature=0.7
    )
    return intro_content.choices[0].message.content

def generate_table_of_contents(book_details):
    toc = f"# Table of Contents\n\nIntroduction\n\n"
    for i, chapter in enumerate(book_details['chapters'], 1):
        toc += f"{i}. {chapter['name']}\n"
    return toc

def save_book_formats(content, title):
    # Save as TXT
    with open(f"{title}.txt", "w", encoding="utf-8") as f:
        f.write(content)

    # Save as DOCX
    doc = Document()
    doc.add_heading(title, 0)
    for paragraph in content.split('\n'):
        doc.add_paragraph(paragraph)
    doc.save(f"{title}.docx")

    # Save as PDF using reportlab
    pdf_filename = f"{title}.pdf"
    doc = SimpleDocTemplate(pdf_filename, pagesize=letter)
    styles = getSampleStyleSheet()
    flowables = []

    for paragraph in content.split('\n'):
        p = Paragraph(paragraph, styles['Normal'])
        flowables.append(p)

    doc.build(flowables)

def extract_text_from_file(file):
    file_extension = os.path.splitext(file.name)[1].lower()
    
    if file_extension == '.pdf':
        pdf_reader = PyPDF2.PdfReader(file)
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text()
    elif file_extension == '.txt':
        text = file.getvalue().decode('utf-8')
    elif file_extension in ['.doc', '.docx']:
        text = docx2txt.process(file)
    else:
        raise ValueError("Unsupported file format")
    
    return text


# World-Building Assistant
def world_building_assistant():
    st.header("World-Building Assistant")
    st.write("Develop and maintain consistent rules, cultures, and environments for your fictional world with AI assistance.")

    if 'world_elements' not in st.session_state:
        st.session_state.world_elements = {}

    # Input for world elements
    element_category = st.selectbox("Element Category", ["Geography", "Culture", "Magic System", "Technology", "History", "Politics"])

    # Custom inputs for each category
    if element_category == "Geography":
        element_name = st.text_input("Location Name")
        climate = st.selectbox("Climate", ["Tropical", "Temperate", "Arctic", "Desert", "Mediterranean", "Alpine", "Subarctic", "Subtropical", "Oceanic", "Semi-arid", "Humid continental", "Tundra"])
        terrain = st.multiselect("Terrain Features", ["Mountains", "Forests", "Rivers", "Oceans", "Plains", "Islands", "Canyons", "Valleys", "Plateaus", "Swamps", "Tundra", "Glaciers", "Volcanoes", "Coral reefs", "Caves", "Fjords", "Sand dunes", "Geysers", "Hot springs"])
        prompt = f"Create a detailed description for a {climate} region named {element_name} featuring {', '.join(terrain)}."

    elif element_category == "Culture":
        element_name = st.text_input("Culture Name")
        social_structure = st.selectbox("Social Structure", ["Hierarchical", "Egalitarian", "Clan-based", "Caste System", "Matriarchal", "Patriarchal", "Meritocracy", "Gerontocracy", "Plutocracy", "Technocracy", "Nomadic", "Tribal confederation"])
        values = st.multiselect("Core Values", ["Honor", "Knowledge", "Nature", "Technology", "Spirituality", "Warfare", "Artistry", "Commerce", "Exploration", "Tradition", "Innovation", "Harmony", "Individualism", "Collectivism", "Asceticism", "Hedonism", "Pacifism", "Mysticism"])
        prompt = f"Describe the {social_structure} culture of {element_name}, emphasizing their focus on {', '.join(values)}."

    elif element_category == "Magic System":
        element_name = st.text_input("Magic System Name")
        source = st.selectbox("Source of Magic", ["Natural Elements", "Divine", "Inner Energy", "Artifacts", "Otherworldly", "Ley Lines", "Celestial Bodies", "Ancient Knowledge", "Emotions", "Life Force", "Dreams", "Music", "Mathematical Formulas", "Blood", "Chaos", "Order"])
        limitations = st.multiselect("Limitations", ["Physical toll", "Rare resources", "Specific bloodline", "Years of study", "Unpredictable", "Time-bound", "Location-dependent", "Emotional state", "Sacrifice required", "Limited uses per day", "Adverse side effects", "Requires multiple practitioners", "Degrades user's lifespan", "Corrupts the user", "Attracts dangerous entities"])
        prompt = f"Detail the {source}-based magic system called {element_name}, including its {', '.join(limitations)} as limitations."

    elif element_category == "Technology":
        element_name = st.text_input("Technology Name")
        tech_level = st.slider("Technology Level", 1, 10, 5)
        focus = st.multiselect("Technological Focus", ["Energy", "Transportation", "Communication", "Warfare", "Medicine", "AI", "Biotechnology", "Nanotechnology", "Space Exploration", "Virtual Reality", "Cybernetics", "Robotics", "Environmental Engineering", "Quantum Computing", "Time Manipulation", "Teleportation", "Genetic Engineering", "Holographics", "Terraforming"])
        prompt = f"Describe the level {tech_level} technology {element_name}, focusing on advancements in {', '.join(focus)}."

    elif element_category == "History":
        element_name = st.text_input("Historical Event/Era Name")
        time_frame = st.selectbox("Time Frame", ["Prehistoric", "Ancient", "Classical", "Medieval", "Renaissance", "Early Modern", "Industrial", "Modern", "Information Age", "Near Future", "Far Future", "Post-Apocalyptic"])
        event_type = st.selectbox("Event Type", ["War", "Discovery", "Cultural Revolution", "Natural Disaster", "Technological Breakthrough", "Religious Movement", "Political Upheaval", "Economic Transformation", "Artistic Renaissance", "Plague/Pandemic", "First Contact", "Environmental Shift", "Mass Migration", "Scientific Revolution", "Golden Age", "Dark Age"])
        prompt = f"Narrate the {time_frame} {event_type} known as {element_name} and its impact on the world."

    elif element_category == "Politics":
        element_name = st.text_input("Political System/Faction Name")
        gov_type = st.selectbox("Government Type", ["Monarchy", "Democracy", "Oligarchy", "Theocracy", "Anarchy", "Republic", "Dictatorship", "Federation", "Communism", "Socialism", "Corporatocracy", "Magocracy", "AI Governance", "Hive Mind", "Tribal Council", "Technocracy", "Direct Democracy", "Stratocracy"])
        key_issues = st.multiselect("Key Political Issues", ["Resource scarcity", "Technological regulation", "Civil rights", "Environmental concerns", "Foreign relations", "Economic inequality", "Education reform", "Healthcare access", "Military expansion", "Magical regulation", "Interspecies relations", "Space colonization", "Artificial intelligence rights", "Genetic modification ethics", "Time travel legislation", "Dimensional boundary control", "Supernatural entity integration"])
        prompt = f"Explain the {gov_type} political system of {element_name}, addressing their stance on {', '.join(key_issues)}."

    if st.button("Generate World Element"):
        if element_name:
            with st.spinner("Generating world element..."):
                response = together_client.chat.completions.create(
                    model="meta-llama/Meta-Llama-3.1-70B-Instruct-Turbo",
                    messages=[
                        {"role": "system", "content": "You are an expert world-building assistant for fiction writers."},
                        {"role": "user", "content": prompt}
                    ],
                    max_tokens=5000,
                    temperature=0.7
                )
                element_description = response.choices[0].message.content

            if element_category not in st.session_state.world_elements:
                st.session_state.world_elements[element_category] = {}
            st.session_state.world_elements[element_category][element_name] = element_description
            st.success(f"{element_name} added to {element_category} successfully!")

    # Display world elements
    if st.session_state.world_elements:
        st.subheader("Your World")
        for category, elements in st.session_state.world_elements.items():
            st.write(f"**{category}**")
            for name, description in elements.items():
                with st.expander(name):
                    st.write(description)
                    if st.button(f"Regenerate {name}", key=f"regen_{category}_{name}"):
                        with st.spinner(f"Regenerating {name}..."):
                            response = together_client.chat.completions.create(
                                model="meta-llama/Meta-Llama-3.1-70B-Instruct-Turbo",
                                messages=[
                                    {"role": "system", "content": "You are an expert world-building assistant for fiction writers."},
                                    {"role": "user", "content": f"Rewrite and improve the following world element description for {category}: {description}"}
                                ],
                                max_tokens=500,
                                temperature=0.7
                            )
                            new_description = response.choices[0].message.content
                            st.session_state.world_elements[category][name] = new_description
                            st.experimental_rerun()

    # AI-powered world consistency check
    if st.button("Check World Consistency"):
        if st.session_state.world_elements:
            all_elements = "\n".join([f"{cat}: {name} - {desc}" for cat, elements in st.session_state.world_elements.items() for name, desc in elements.items()])
            with st.spinner("Analyzing world consistency..."):
                response = together_client.chat.completions.create(
                    model="meta-llama/Meta-Llama-3.1-70B-Instruct-Turbo",
                    messages=[
                        {"role": "system", "content": "You are an expert world-building consultant for fiction writers."},
                        {"role": "user", "content": f"Analyze the following world elements for consistency and provide suggestions for improvement:\n\n{all_elements}"}
                    ],
                    max_tokens=1000,
                    temperature=0.7
                )
                consistency_analysis = response.choices[0].message.content
            st.subheader("World Consistency Analysis")
            st.write(consistency_analysis)
        else:
            st.warning("Add some world elements before checking consistency.")

    # Export world-building data
    if st.session_state.world_elements:
        if st.button("Export World Data"):
            export_data = {category: {name: desc for name, desc in elements.items()} for category, elements in st.session_state.world_elements.items()}
            st.download_button(
                label="Download World Data as JSON",
                data=json.dumps(export_data, indent=2),
                file_name="world_building_data.json",
                mime="application/json"
            )

def get_audio_data():
    CHUNK = 1024
    FORMAT = pyaudio.paInt16
    CHANNELS = 1
    RATE = 44100
    RECORD_SECONDS = 0.1

    p = pyaudio.PyAudio()

    stream = p.open(format=FORMAT,
                    channels=CHANNELS,
                    rate=RATE,
                    input=True,
                    frames_per_buffer=CHUNK)

    frames = []

    for i in range(0, int(RATE / CHUNK * RECORD_SECONDS)):
        data = stream.read(CHUNK)
        frames.append(np.frombuffer(data, dtype=np.int16))

    stream.stop_stream()
    stream.close()
    p.terminate()

    return np.concatenate(frames)

def update_audio_chart(chart, audio_data):
    chart.plotly_chart(go.Figure(data=go.Scatter(y=audio_data, mode='lines')), use_container_width=True)

def create_radar_chart(stats):
    df = pd.DataFrame(dict(
        r=list(stats.values()),
        theta=list(stats.keys())
    ))
    fig = px.line_polar(df, r='r', theta='theta', line_close=True)
    fig.update_traces(fill='toself')
    fig.update_layout(
        polar=dict(
            radialaxis=dict(visible=True, range=[0, 10])
        ),
        showlegend=False
    )
    return fig

import html
from urllib.parse import urlparse
from googleapiclient.discovery import build
from googleapiclient.errors import HttpError
import requests
from bs4 import BeautifulSoup

if "research_data" not in st.session_state:
    st.session_state.research_data = {
        "chats": [],
        "folders": {},
        "ideas": [],
        "rough_thoughts": []
    }

def save_research_data():
    with open("research_data.json", "w") as f:
        json.dump(st.session_state.research_data, f)

def load_research_data():
    if os.path.exists("research_data.json"):
        with open("research_data.json", "r") as f:
            st.session_state.research_data = json.load(f)

def search_web(user_input, api_key="AIzaSyD-1OMuZ0CxGAek0PaXrzHOmcDWFvZQtm8", cse_id="877170db56f5c4629"):
    """
    Perform web search using Google Custom Search API and format results for the AI.
    """
    try:
        # Build the service
        service = build("customsearch", "v1", developerKey=api_key)
        
        # Perform the search
        result = service.cse().list(
            q=user_input,
            cx=cse_id,
            num=5,
            start=1
        ).execute()
        
        # Process and format the results
        processed_results = []
        
        if 'items' in result:
            for i, item in enumerate(result['items'], 1):
                title = html.unescape(item.get('title', ''))
                description = html.unescape(item.get('snippet', ''))
                link = item.get('link', '')
                domain = urlparse(link).netloc
                
                # Fetch and extract main content from the webpage
                try:
                    response = requests.get(link, timeout=5)
                    soup = BeautifulSoup(response.content, 'html.parser')
                    
                    # Extract main content (this is a simple example, you might need to adjust based on the sites you're scraping)
                    main_content = ' '.join([p.text for p in soup.find_all('p')])
                    main_content = main_content[:1000]  # Limit to 1000 characters
                except:
                    main_content = description
                
                # Create a formatted source entry
                source_entry = {
                    "id": i,
                    "title": title,
                    "link": link,
                    "description": description,
                    "domain": domain,
                    "content": f"{main_content}\n\nFrom: {domain}"
                }
                processed_results.append(source_entry)
        
        return processed_results
    
    except HttpError as e:
        print(f"An HTTP error {e.resp.status} occurred: {e.content}")
        return []
    except Exception as e:
        print(f"Error performing search: {str(e)}")
        return []

def generate_ai_response_ws(prompt, sources):
    """
    Generate AI response with proper source integration.
    """
    together_client = Together(api_key=TOGETHER_API_KEY)
    
    # Create a detailed system prompt that enforces source citation
    system_prompt = """You are a helpful AI research assistant. Follow these rules strictly:

1. Every factual claim must be supported by the provided sources using numbered citations [1], [2], etc.
2. Format your response in clear paragraphs with proper markdown
3. For each citation, specify the source title and domain in parentheses
4. Include a "Sources" section at the end listing all referenced sources with their full URLs
5. If you need to provide context beyond the sources, clearly mark it as "Additional Context"
6. Quote important phrases from sources when appropriate, using quotation marks

Example format:

According to [1] (Source Title, domain.com), "quoted text here"...

Sources:
1. Source Title (URL)
2. Source Title (URL)"""

    # Format sources into a structured reference
    source_text = "\n\n".join([
        f"Source {s['id']}:\nTitle: {s['title']}\nURL: {s['link']}\nContent: {s['content']}\nDomain: {s['domain']}"
        for s in sources
    ])
    
    messages = [
        {"role": "system", "content": system_prompt},
        {"role": "user", "content": f"""Question: {prompt}

Available Sources:
{source_text}

Please provide a comprehensive answer using these sources."""}
    ]
    
    response = together_client.chat.completions.create(
        model="meta-llama/Meta-Llama-3.1-70B-Instruct-Turbo",
        messages=messages,
        max_tokens=1500,
        temperature=0.7,
        top_p=0.7,
        top_k=50,
        repetition_penalty=1,
        stop=["<|eot_id|>", "<|eom_id|>"],
        stream=True
    )
    
    full_response = ""
    for token in response:
        if hasattr(token, 'choices'):
            content = token.choices[0].delta.content
            full_response += content
            yield content

def book_research():
    st.title("Advanced Book Research")
    
    tab1, tab2 = st.tabs(["Research Assistant", "Research Organizer"])
    
    with tab1:
        st.header("Research Assistant")
        
        # Chat history
        st.subheader("Chat History")
        for i, chat in enumerate(st.session_state.research_data["chats"]):
            col1, col2 = st.columns([0.9, 0.1])
            with col1:
                st.markdown(chat['content'])  # Changed to markdown for better formatting
            with col2:
                if st.button("Delete", key=f"delete_chat_{i}"):
                    st.session_state.research_data["chats"].pop(i)
                    save_research_data()
                    st.experimental_rerun()
        
        # Clear chat history
        if st.button("Clear Chat History"):
            st.session_state.research_data["chats"] = []
            save_research_data()
            st.success("Chat history cleared!")
            st.experimental_rerun()
        
        # User input
        user_input = st.text_input("Ask a question for your book research:")
        if st.button("Submit"):
            if user_input:
                # Add user message to chat history with formatting
                st.session_state.research_data["chats"].append({
                    "role": "user", 
                    "content": f"**Question:** {user_input}"
                })
                
                with st.spinner("Searching and analyzing..."):
                    # Search the web
                    search_results = search_web(user_input)
                    
                    if search_results:
                        # Generate AI response
                        st.write("AI Response:")
                        response_placeholder = st.empty()
                        full_response = ""
                        
                        # Pass the full search_results to generate_ai_response_ws
                        for token in generate_ai_response_ws(user_input, search_results):
                            full_response += token
                            response_placeholder.markdown(full_response)
                        
                        # Add AI response to chat history
                        st.session_state.research_data["chats"].append({
                            "role": "assistant",
                            "content": full_response
                        })
                        
                        # Display sources in a collapsible section
                        with st.expander("View Sources", expanded=True):
                            st.markdown("### Referenced Sources")
                            for result in search_results:
                                st.markdown(
                                    f"""---
**[{result['id']}] {result['title']}**  
📍 Domain: {result['domain']}  
🔗 [Link]({result['link']})  
📝 Summary: {result['description']}
                                    """)
                        
                        # Save research data
                        save_research_data()
                    else:
                        st.error("No results found. Please try a different search query.")
    
    with tab2:
        st.header("Research Organizer")
        
        col1, col2 = st.columns(2)
        
        with col1:
            st.subheader("Folders")
            
            # Create new folder
            new_folder = st.text_input("Create a new folder:")
            if st.button("Create Folder"):
                if new_folder and new_folder not in st.session_state.research_data["folders"]:
                    st.session_state.research_data["folders"][new_folder] = []
                    save_research_data()
                    st.success(f"Folder '{new_folder}' created successfully!")
            
            # Display folders and their contents
            for folder, contents in st.session_state.research_data["folders"].items():
                with st.expander(folder):
                    for i, item in enumerate(contents):
                        col1, col2 = st.columns([0.9, 0.1])
                        with col1:
                            st.write(item)
                        with col2:
                            if st.button("Delete", key=f"delete_folder_item_{folder}_{i}"):
                                st.session_state.research_data["folders"][folder].pop(i)
                                save_research_data()
                                st.experimental_rerun()
                    
                    # Add new item to folder
                    new_item = st.text_area(f"Add new item to {folder}:")
                    if st.button(f"Add to {folder}"):
                        if new_item:
                            st.session_state.research_data["folders"][folder].append(new_item)
                            save_research_data()
                            st.success(f"Item added to '{folder}' successfully!")
                            st.experimental_rerun()
        
        with col2:
            st.subheader("Quick Ideas")
            
            # Add new idea
            new_idea = st.text_area("Add a new idea:")
            if st.button("Save Idea"):
                if new_idea:
                    st.session_state.research_data["ideas"].append({
                        "content": new_idea,
                        "timestamp": datetime.now().strftime("%Y-%m-%d %H:%M:%S")
                    })
                    save_research_data()
                    st.success("Idea saved successfully!")
                    st.experimental_rerun()
            
            # Display ideas
            for i, idea in enumerate(st.session_state.research_data["ideas"]):
                col1, col2 = st.columns([0.9, 0.1])
                with col1:
                    st.text(f"{idea['timestamp']}: {idea['content']}")
                with col2:
                    if st.button("Delete", key=f"delete_idea_{i}"):
                        st.session_state.research_data["ideas"].pop(i)
                        save_research_data()
                        st.experimental_rerun()
        
        # Rough Thoughts
        st.subheader("Rough Thoughts")
        new_thought = st.text_area("Jot down a rough thought:")
        if st.button("Save Thought"):
            if new_thought:
                st.session_state.research_data["rough_thoughts"].append({
                    "content": new_thought,
                    "timestamp": datetime.now().strftime("%Y-%m-%d %H:%M:%S")
                })
                save_research_data()
                st.success("Rough thought saved successfully!")
                st.experimental_rerun()
        
        # Display rough thoughts
        for i, thought in enumerate(st.session_state.research_data["rough_thoughts"]):
            col1, col2 = st.columns([0.9, 0.1])
            with col1:
                st.text(f"{thought['timestamp']}: {thought['content']}")
            with col2:
                if st.button("Delete", key=f"delete_thought_{i}"):
                    st.session_state.research_data["rough_thoughts"].pop(i)
                    save_research_data()
                    st.experimental_rerun()
        
        # Clear all data
        if st.button("Clear All Research Data"):
            st.session_state.research_data = {
                "chats": [],
                "folders": {},
                "ideas": [],
                "rough_thoughts": []
            }
            save_research_data()
            st.success("All research data cleared!")
            st.experimental_rerun()
        
        # Export research data
        if st.button("Export Research Data"):
            st.download_button(
                label="Download Research Data",
                data=json.dumps(st.session_state.research_data, indent=2),
                file_name="research_data_export.json",
                mime="application/json"
            )

def script_generator():
    st.header("AI Script Generator")
    
    if 'demo_data' not in st.session_state:
        st.session_state.demo_data = None
    
    if st.button("Load Romance Movie Demo"):
        st.session_state.demo_data = {
            "title": "Love in the City",
            "genre": "Romance",
            "script_type": "Movie",
            "setting": "New York City",
            "theme": "Finding love in unexpected places",
            "tone": "Heartwarming and humorous",
            "target_audience": "Adults",
            "duration": 120,
            "logline": "A workaholic lawyer and a free-spirited artist clash and connect as they navigate the bustling streets of New York.",
            "pretext": "In a world where dating apps rule and genuine connections seem rare.",
            "num_acts": 3,
            "num_characters": 3,
            "characters": [
                {"name": "Emma Thompson", "role": "Workaholic Lawyer", "description": "Ambitious, organized, and initially closed off to romance"},
                {"name": "Jack Reynolds", "role": "Free-spirited Artist", "description": "Charismatic, spontaneous, and always looking for inspiration"},
                {"name": "Sarah Chen", "role": "Emma's Best Friend", "description": "Supportive, outgoing, and always trying to set Emma up"}
            ],
            "acts": [
                {
                    "description": "Set up the characters and their conflicting worlds",
                    "num_scenes": 3,
                    "scenes": [
                        {
                            "description": "Emma rushing through morning routine",
                            "characters": ["Emma Thompson"],
                            "location": "Emma's Apartment"
                        },
                        {
                            "description": "Jack painting in the park",
                            "characters": ["Jack Reynolds"],
                            "location": "Central Park"
                        },
                        {
                            "description": "First encounter - coffee spill",
                            "characters": ["Emma Thompson", "Jack Reynolds"],
                            "location": "Busy Street"
                        }
                    ]
                }
            ],
            "script_style": "Standard"
        }
        st.rerun()

    def get_demo_value(key, default=""):
        return st.session_state.demo_data.get(key, default) if st.session_state.demo_data else default

    script_type = st.selectbox("Select Script Type", 
                              ["Movie", "Theater Play", "Web Series", "TV Show"],
                              index=0 if not st.session_state.demo_data else ["Movie", "Theater Play", "Web Series", "TV Show"].index(get_demo_value("script_type", "Movie")))
    
    genres = ["Action", "Comedy", "Drama", "Sci-Fi", "Horror", "Romance", "Thriller", "Fantasy", "Mystery", "Adventure"]
    genre = st.selectbox("Select Genre", genres, index=genres.index(get_demo_value("genre", "Action")))
    
    col1, col2 = st.columns(2)
    with col1:
        title = st.text_input("Script Title", value=get_demo_value("title"))
        setting = st.text_input("Setting", value=get_demo_value("setting"))
        target_audience = st.selectbox("Target Audience", 
                                     ["General", "Children", "Teenagers", "Adults", "Mature"],
                                     index=["General", "Children", "Teenagers", "Adults", "Mature"].index(get_demo_value("target_audience", "General")))
    with col2:
        theme = st.text_input("Theme", value=get_demo_value("theme"))
        tone = st.text_input("Tone", value=get_demo_value("tone"))
        duration = st.slider("Approximate Duration (minutes)", 15, 180, value=int(get_demo_value("duration", "90")))

    logline = st.text_area("Logline", value=get_demo_value("logline"), height=50)
    pretext = st.text_area("Pretext", value=get_demo_value("pretext"), height=100)
    
    num_acts = st.number_input("Number of Acts", 1, 5, value=int(get_demo_value("num_acts", "3")))
    
    st.subheader("Character Setup")
    num_characters = st.number_input("Number of Main Characters", 1, 10, value=int(get_demo_value("num_characters", "3")))
    
    characters = []
    demo_characters = get_demo_value("characters", [])
    for i in range(num_characters):
        col1, col2, col3 = st.columns(3)
        with col1:
            name = st.text_input(f"Character {i+1} Name", 
                               value=demo_characters[i]["name"] if i < len(demo_characters) else "")
        with col2:
            role = st.text_input(f"Character {i+1} Role",
                               value=demo_characters[i]["role"] if i < len(demo_characters) else "")
        with col3:
            description = st.text_area(f"Character {i+1} Description",
                                     value=demo_characters[i]["description"] if i < len(demo_characters) else "",
                                     height=50)
        characters.append({"name": name, "role": role, "description": description})

    acts = []
    demo_acts = get_demo_value("acts", [])
    for act in range(num_acts):
        st.subheader(f"Act {act + 1}")
        act_description = st.text_area(
            f"Act {act + 1} Description",
            value=demo_acts[act]["description"] if act < len(demo_acts) else "",
            height=50,
            key=f"act_desc_{act}"
        )
        num_scenes = st.number_input(
            f"Number of Scenes in Act {act + 1}", 
            1, 10, 
            value=int(demo_acts[act]["num_scenes"]) if act < len(demo_acts) else 3,
            key=f"num_scenes_{act}"
        )
        
        scenes = []
        demo_scenes = demo_acts[act]["scenes"] if act < len(demo_acts) else []
        for scene in range(num_scenes):
            st.markdown(f"**Scene {scene + 1}**")
            scene_description = st.text_input(
                f"Description: Act {act + 1}, Scene {scene + 1}",
                value=demo_scenes[scene]["description"] if scene < len(demo_scenes) else "",
                key=f"scene_desc_{act}_{scene}"
            )
            scene_characters = st.multiselect(
                f"Characters in Scene {scene + 1}",
                [char['name'] for char in characters],
                default=demo_scenes[scene]["characters"] if scene < len(demo_scenes) else [],
                key=f"scene_chars_{act}_{scene}"  # Added unique key here
            )
            scene_location = st.text_input(
                f"Location: Act {act + 1}, Scene {scene + 1}",
                value=demo_scenes[scene]["location"] if scene < len(demo_scenes) else "",
                key=f"scene_loc_{act}_{scene}"
            )
            scenes.append({
                "description": scene_description,
                "characters": scene_characters,
                "location": scene_location
            })
        acts.append({
            "description": act_description,
            "num_scenes": num_scenes,
            "scenes": scenes
        })

    script_style = st.selectbox(
        "Script Style",
        ["Standard", "Tarantino-esque", "Sorkin-style Dialogue", "Nolan-like Complexity"],
        index=["Standard", "Tarantino-esque", "Sorkin-style Dialogue", "Nolan-like Complexity"].index(get_demo_value("script_style", "Standard"))
    )

    if st.button("Generate Script"):
        with st.spinner("Generating your script... This may take a few minutes."):
            full_script = generate_full_script(title, genre, setting, theme, tone, target_audience, logline, pretext, characters, acts, script_style, script_type)
            
            st.subheader("Generated Script")
            st.text_area("Script Content", full_script, height=600)
            
            st.subheader("Download Options")
            col1, col2, col3 = st.columns(3)
            
            with col1:
                if st.button("Download as TXT"):
                    download_as_txt(full_script, title)
            with col2:
                if st.button("Download as DOCX"):
                    download_as_docx(full_script, title)
            with col3:
                if st.button("Download as PDF"):
                    download_as_pdf(full_script, title)

            st.subheader("Edit Script")
            edited_script = st.text_area("Edit your script here", full_script, height=600)
            if st.button("Save Edited Script"):
                save_edited_script(edited_script, title)

def generate_full_script(title, genre, setting, theme, tone, target_audience, logline, pretext, characters, acts, script_style, script_type):
    full_script = f"Title: {title}\n\nGenre: {genre}\n\nSetting: {setting}\n\nTheme: {theme}\n\nTone: {tone}\n\nLogline: {logline}\n\n"
    full_script += "Characters:\n"
    for char in characters:
        full_script += f"- {char['name']} ({char['role']}): {char['description']}\n"
    full_script += "\n"

    for act_num, act in enumerate(acts, 1):
        full_script += f"\nACT {act_num}: {act['description']}\n\n"
        for scene_num, scene in enumerate(act['scenes'], 1):
            # Prepare the prompt for each scene
            prompt = f"Generate a scene for a {script_type} script titled '{title}' in the {genre} genre. "
            prompt += f"This is Act {act_num}, Scene {scene_num}. "
            prompt += f"Setting: {setting}. Theme: {theme}. Tone: {tone}. "
            prompt += f"Target audience: {target_audience}. Logline: {logline}. "
            prompt += f"Scene description: {scene['description']}. "
            prompt += f"Characters in this scene: {', '.join(scene['characters'])}. "
            prompt += f"Scene location: {scene['location']}. "
            if pretext:
                prompt += f"Background information: {pretext}. "
            prompt += f"Write the scene in a {script_style} style, following standard script formatting conventions."

            # Generate the scene using Together AI
            response = together_client.chat.completions.create(
                model="meta-llama/Meta-Llama-3.1-70B-Instruct-Turbo",
                messages=[
                    {"role": "system", "content": f"You are an expert {script_type} scriptwriter specializing in {genre} scripts with a {script_style} style."},
                    {"role": "user", "content": prompt}
                ],
                max_tokens=6048,  # Adjust as needed
                temperature=0.7
            )
            
            scene_content = response.choices[0].message.content
            full_script += f"Scene {scene_num} - {scene['location']}\n\n{scene_content}\n\n"

    return full_script

def download_as_txt(script, title):
    tmp = io.BytesIO()
    tmp.write(script.encode())
    tmp.seek(0)
    st.download_button(
        label="Download TXT",
        data=tmp,
        file_name=f"{title}_script.txt",
        mime="text/plain"
    )

def download_as_docx(script, title):
    doc = Document()
    doc.add_heading(title, 0)
    doc.add_paragraph(script)
    docx_file = io.BytesIO()
    doc.save(docx_file)
    docx_file.seek(0)
    st.download_button(
        label="Download DOCX",
        data=docx_file,
        file_name=f"{title}_script.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

def download_as_pdf(script, title):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    pdf.multi_cell(0, 10, script)
    pdf_file = io.BytesIO()
    pdf.output(pdf_file)
    pdf_file.seek(0)
    st.download_button(
        label="Download PDF",
        data=pdf_file,
        file_name=f"{title}_script.pdf",
        mime="application/pdf"
    )

def save_edited_script(script, title):
    # Here you can add logic to save the edited script, e.g., to a database or file
    st.success(f"Edited script '{title}' saved successfully!")

def load_romance_movie_demo():
    demo_data = {
        "title": "Love in the City",
        "genre": "Romance",
        "setting": "New York City",
        "theme": "Finding love in unexpected places",
        "tone": "Heartwarming and humorous",
        "target_audience": "Adults",
        "duration": 120,
        "logline": "A workaholic lawyer and a free-spirited artist clash and connect as they navigate the bustling streets of New York, learning to balance their careers with matters of the heart.",
        "pretext": "In a world where dating apps rule and genuine connections seem rare, two polar opposites find that sometimes love is just around the corner – literally.",
        "characters": [
            {"name": "Emma Thompson", "role": "Workaholic Lawyer", "description": "Ambitious, organized, and initially closed off to romance"},
            {"name": "Jack Reynolds", "role": "Free-spirited Artist", "description": "Charismatic, spontaneous, and always looking for inspiration"},
            {"name": "Sarah Chen", "role": "Emma's Best Friend", "description": "Supportive, outgoing, and always trying to set Emma up"}
        ],
        "acts": [
            {
                "description": "Set up the characters and their conflicting worlds",
                "num_scenes": 3,
                "scenes": [
                    {
                        "description": "Emma rushing through morning routine, barely making it to important meeting",
                        "characters": ["Emma Thompson"],
                        "location": "Emma's Apartment / Law Firm"
                    },
                    {
                        "description": "Jack painting in the park, interacting with various New Yorkers",
                        "characters": ["Jack Reynolds"],
                        "location": "Central Park"
                    },
                    {
                        "description": "Emma and Jack's first encounter - a coffee spill incident",
                        "characters": ["Emma Thompson", "Jack Reynolds"],
                        "location": "Busy New York Street"
                    }
                ]
            },
            {
                "description": "Developing relationship and overcoming obstacles",
                "num_scenes": 3,
                "scenes": [
                    {
                        "description": "Emma and Jack keep running into each other, slowly warming up",
                        "characters": ["Emma Thompson", "Jack Reynolds"],
                        "location": "Various New York locations"
                    },
                    {
                        "description": "Emma's work crisis clashes with Jack's art show opening",
                        "characters": ["Emma Thompson", "Jack Reynolds", "Sarah Chen"],
                        "location": "Law Firm / Art Gallery"
                    },
                    {
                        "description": "Emma and Jack have a heart-to-heart, realizing their feelings",
                        "characters": ["Emma Thompson", "Jack Reynolds"],
                        "location": "Brooklyn Bridge at sunset"
                    }
                ]
            },
            {
                "description": "Resolution and coming together",
                "num_scenes": 3,
                "scenes": [
                    {
                        "description": "Emma makes a grand gesture at Jack's exhibition",
                        "characters": ["Emma Thompson", "Jack Reynolds", "Sarah Chen"],
                        "location": "Major Art Gallery"
                    },
                    {
                        "description": "Jack surprises Emma at her office with a romantic gesture",
                        "characters": ["Emma Thompson", "Jack Reynolds"],
                        "location": "Law Firm"
                    },
                    {
                        "description": "Emma and Jack's new life together, balancing work and love",
                        "characters": ["Emma Thompson", "Jack Reynolds"],
                        "location": "New York City streets and their shared apartment"
                    }
                ]
            }
        ],
        "script_style": "Standard"
    }

    # Update all the input fields with the demo data
    for key, value in demo_data.items():
        if key == 'characters':
            st.session_state['num_characters'] = len(value)
            for i, char in enumerate(value):
                st.session_state[f'character_{i+1}_name'] = char['name']
                st.session_state[f'character_{i+1}_role'] = char['role']
                st.session_state[f'character_{i+1}_description'] = char['description']
        elif key == 'acts':
            st.session_state['num_acts'] = len(value)
            for i, act in enumerate(value):
                st.session_state[f'act_{i+1}_description'] = act['description']
                st.session_state[f'act_{i+1}_num_scenes'] = act['num_scenes']
                for j, scene in enumerate(act['scenes']):
                    st.session_state[f'act_{i+1}_scene_{j+1}_description'] = scene['description']
                    st.session_state[f'act_{i+1}_scene_{j+1}_characters'] = scene['characters']
                    st.session_state[f'act_{i+1}_scene_{j+1}_location'] = scene['location']
        else:
            st.session_state[key] = value

    st.success("Romance Movie Demo loaded successfully!")

def content_generator():
    st.title("Content Generator")
    
    content_types = [
        "Content Writing", "Advertising Writing", "Marketing Writing",
        "Grant Writing", "Proposal Writing"
    ]
    
    selected_type = st.selectbox("Select Content Type", content_types)
    
    text_content_generator(selected_type)

def text_content_generator(content_type):
    st.subheader(f"{content_type}")
    
    # Input fields
    topic = st.text_input("Topic or Title")
    keywords = st.text_input("Enter Keywords (comma-separated)")
    tone = st.selectbox("Tone", ["Professional", "Casual", "Humorous", "Formal", "Persuasive"])
    target_audience = st.text_input("Target Audience")
    word_count = st.slider("Word Count", 100, 2000, 500)
    
    additional_instructions = st.text_area("Additional Instructions (Optional)")
    
    if st.button("Generate Content"):
        if not topic:
            st.warning("Please enter a topic.")
            return
        
        keyword_list = [k.strip() for k in keywords.split(',') if k.strip()]
        
        prompt = f"""Generate {content_type} content with the following details:
        Topic: {topic}
        Keywords: {', '.join(keyword_list)}
        Tone: {tone}
        Target Audience: {target_audience}
        Word Count: Approximately {word_count} words
        Additional Instructions: {additional_instructions}

        Please provide well-structured, engaging content that meets the specified requirements."""
        
        with st.spinner("Generating content..."):
            response = together_client.chat.completions.create(
                model="meta-llama/Meta-Llama-3.1-70B-Instruct-Turbo",
                messages=[
                    {"role": "system", "content": f"You are an expert in {content_type}."},
                    {"role": "user", "content": prompt}
                ],
                max_tokens=word_count * 2,  # Assuming an average of 2 tokens per word
                temperature=0.7
            )
            
            generated_content = response.choices[0].message.content
            
            st.subheader("Generated Content")
            st.write(generated_content)
            
            # Download options
            st.download_button(
                label="Download as TXT",
                data=generated_content,
                file_name=f"{content_type.lower().replace(' ', '_')}.txt",
                mime="text/plain"
            )

def main():
    st.set_page_config(page_title="WORDCRAFT - AI Automation for Writers", page_icon="📚", layout="wide")

    menu = ["Home", "Create Character", "Chat with Characters", "Generate Book", "Convert to Audiobook",  "AI Script Generator", "Content Generator", "Advanced Book Research",
            "Book Outline Generator", "Character Development Workshop", "Writing Prompts Generator", "World-Building Assistant", "Interactive Character Board"]
    choice = st.sidebar.selectbox("Menu", menu)

    if choice == "Home":
        st.title("WORDCRAFT - AI Automation for Writers")
        st.write("Welcome to WORDCRAFT, your AI-powered writing assistant!")
        
        st.header("Unlock Your Writing Potential with AI")
        st.write("""
        WORDCRAFT is designed to revolutionize your writing process and help you generate thousands of dollars per month through KDP self-publishing. Our AI-powered tools streamline your workflow, boost creativity, and enhance your productivity.
        """)
        
        st.subheader("Key Features:")
        col1, col2 = st.columns(2)
        
        with col1:
            st.markdown("""
            - **AI Character Creation**: Bring your characters to life with unique voices and personalities.
            - **Interactive Character Chats**: Develop your characters through dynamic conversations.
            - **Automated Book Generation**: Create full-length books with AI assistance.
            - **Audiobook Conversion**: Transform your text into professional audiobooks.
            - **Advanced Book Research**: Book research powered with web and AI.
            """)
        
        with col2:
            st.markdown("""
            - **Book Outline Generator**: Craft detailed outlines for your next bestseller.
            - **Character Development Workshop**: Deep dive into character creation and evolution.
            - **Writing Prompts Generator**: Spark your creativity with AI-generated prompts.
            - **Multi-format Export**: Save your work in TXT, DOCX, and PDF formats.
            - **Script and Content Generator**: Generates scripts and content for various purposes.
            """)
        
        st.header("How WORDCRAFT Boosts Your KDP Earnings")
        st.write("""
        1. **Rapid Content Creation**: Generate high-quality books faster than ever before.
        2. **Diverse Genre Expertise**: Our AI adapts to any genre, helping you tap into lucrative markets.
        3. **Consistent Output**: Maintain a steady publishing schedule to build your author brand.
        4. **Enhanced Quality**: AI-assisted editing and character development improve your storytelling.
        5. **Audiobook Integration**: Easily create audiobooks to diversify your income streams.
        """)
        
        st.header("Get Started Today!")
        st.write("""
        Explore our features using the sidebar menu and start transforming your writing career. With WORDCRAFT, you're not just writing—you're crafting your path to financial success in the world of self-publishing.
        """)
    
    elif choice == "Create Character":
        st.header("Create a New Character")
        col1, col2 = st.columns(2)

        with col1:
            new_name = st.text_input("Character Name")
            new_description = st.text_area("Short Character Description (500 chars)", max_chars=500)

        with col2:
            new_detailed_description = st.text_area("Detailed Character Description (Optional)")

        if st.button("Create Character") and new_name and new_description:
            st.session_state.characters[new_name] = {
                "description": new_description,
                "detailed_description": new_detailed_description
            }
            st.success(f"Character '{new_name}' created successfully!")

    elif choice == "Chat with Characters":
        st.header("Chat with Characters")

        if 'messages' not in st.session_state:
            st.session_state.messages = []

        col1, col2 = st.columns([1, 3])

        with col1:
            st.subheader("Characters")
            for character in st.session_state.characters:
                if st.button(character, key=f"btn_{character}"):
                    st.session_state.current_character = character
                    st.session_state.messages = []

        with col2:
            if 'current_character' in st.session_state and st.session_state.current_character in st.session_state.characters:
                current_char = st.session_state.characters[st.session_state.current_character]
                st.subheader(f"Chatting with {st.session_state.current_character}")

                # Display chat history
                for message in st.session_state.messages:
                    with st.chat_message(message["role"]):
                        st.write(message["content"])

                language = st.selectbox("Select Language", ["English", "Spanish", "French", "German", "Chinese", "Japanese"])
                user_input = st.chat_input("Type your message here...")

                if user_input:
                    # Add user message to chat history
                    st.session_state.messages.append({"role": "user", "content": user_input})

                    # Display user message
                    with st.chat_message("user"):
                        st.write(user_input)

                    with st.spinner("Generating response..."):
                        ai_response = generate_ai_response(st.session_state.current_character, user_input, language)

                    # Add AI response to chat history
                    st.session_state.messages.append({
                        "role": "assistant", 
                        "content": ai_response
                    })

                    # Display AI response
                    with st.chat_message("assistant"):
                        st.write(ai_response)

                    st.experimental_rerun()

            else:
                st.info("Please select a character to start chatting.")
    elif choice == "Generate Book":
        st.header("Generate a Book")
        
        genre_type = st.radio("Choose genre type:", ("Fiction", "Non-fiction"))
        
        if genre_type == "Fiction":
            genre = st.selectbox("Select Fiction Genre", fiction_genres)
        else:
            genre = st.selectbox("Select Non-fiction Genre", non_fiction_genres)
        
        title = st.text_input("Book Title")
        author = st.text_input("Author Name")
        description = st.text_area("Book Description")
        
        # New section for writing style
        st.subheader("Writing Style")
        writing_style = st.selectbox("Choose Writing Style", [
            "Custom",
            "Ernest Hemingway",
            "Jane Austen",
            "Stephen King",
            "J.K. Rowling",
            "George R.R. Martin",
            "Agatha Christie"
        ])
        
        if writing_style == "Custom":
            style_sample = st.text_area("Provide a sample of the desired writing style (1000 characters max):", max_chars=1000)
        else:
            style_sample = f"Write in the style of {writing_style}."
        
        num_chapters = st.number_input("Number of Chapters", min_value=1, value=5)
        
        chapters = []
        for i in range(num_chapters):
            st.subheader(f"Chapter {i+1}")
            chapter_name = st.text_input(f"Chapter {i+1} Name")
            if genre_type == "Fiction":
                num_scenes = st.number_input(f"Number of Scenes in Chapter {i+1}", min_value=1, value=3)
                scenes = []
                for j in range(num_scenes):
                    scene_description = st.text_area(f"Scene {j+1} Description (Chapter {i+1})")
                    scenes.append(scene_description)
                chapters.append({"name": chapter_name, "scenes": scenes})
            else:
                chapter_description = st.text_area(f"Chapter {i+1} Description")
                num_parts = st.number_input(f"Number of Parts in Chapter {i+1}", min_value=1, value=3)
                parts = []
                for j in range(num_parts):
                    part_description = st.text_area(f"Part {j+1} Description (Chapter {i+1})")
                    parts.append(part_description)
                chapters.append({"name": chapter_name, "description": chapter_description, "parts": parts})
        
        if st.button("Generate Book"):
            book_details = {
                "genre_type": genre_type,
                "genre": genre,
                "title": title,
                "author": author,
                "description": description,
                "writing_style": writing_style,
                "style_sample": style_sample,
                "chapters": chapters
            }
            
            st.session_state.book_details = book_details
            st.session_state.generated_book = ""
            
            with st.spinner("Generating book introduction..."):
                introduction = generate_book_introduction(book_details)
                st.session_state.generated_book += f"# {title}\nBy {author}\n\n{introduction}\n\n"
            
            toc = generate_table_of_contents(book_details)
            st.session_state.generated_book += f"{toc}\n\n"
            
            for i, chapter in enumerate(chapters):
                with st.spinner(f"Generating Chapter {i+1}: {chapter['name']}..."):
                    chapter_content = generate_book_chapter(book_details, i)
                    st.session_state.generated_book += f"# Chapter {i+1}: {chapter['name']}\n\n{chapter_content}\n\n"
                st.success(f"Chapter {i+1} generated successfully!")
            
            st.success("Book generated successfully!")
            st.text_area("Generated Book", st.session_state.generated_book, height=300)
            
            save_book_formats(st.session_state.generated_book, title)
            
            st.download_button(
                label="Download as TXT",
                data=st.session_state.generated_book,
                file_name=f"{title}.txt",
                mime="text/plain"
            )
            
            with open(f"{title}.docx", "rb") as docx_file:
                st.download_button(
                    label="Download as DOCX",
                    data=docx_file,
                    file_name=f"{title}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
            
            with open(f"{title}.pdf", "rb") as pdf_file:
                st.download_button(
                    label="Download as PDF",
                    data=pdf_file,
                    file_name=f"{title}.pdf",
                    mime="application/pdf"
                )

    elif choice == "Convert to Audiobook":
        st.header("Convert to Audiobook")
        st.write("Transform your written work into an audiobook using text-to-speech.")
    
        col1, col2 = st.columns([2, 1])
    
        with col1:
            upload_type = st.radio("Choose input method:", ("Upload File", "Paste Text"), horizontal=True)
    
            if upload_type == "Upload File":
                uploaded_file = st.file_uploader("Upload your book file", type=['txt', 'pdf', 'doc', 'docx'])
                if uploaded_file:
                    with st.spinner("Extracting text from file..."):
                        text = extract_text_from_file(uploaded_file)
                    st.success("Text extracted successfully!")
            else:
                text = st.text_area("Enter your book text here", height=300)
    
            if 'text' in locals() and text:
                st.write("Preview:")
                st.info(text[:500] + "..." if len(text) > 500 else text)
    
        with col2:
            st.subheader("Voice Settings")
    
            # Get available voices from pyttsx3
            voices = engine.getProperty('voices')
            voice_options = {voice.name: i for i, voice in enumerate(voices)}
    
            selected_voice = st.selectbox("Select a voice", list(voice_options.keys()))
            speed = st.slider("Speech Rate", min_value=100, max_value=300, value=200, step=20)
            volume = st.slider("Volume", min_value=0.0, max_value=1.0, value=1.0, step=0.1)
    
        if st.button("Convert to Audiobook", type="primary"):
            if 'text' in locals() and text:
                try:
                    with st.spinner("Converting text to speech..."):
                        # Configure pyttsx3 settings
                        engine.setProperty('rate', speed)
                        engine.setProperty('volume', volume)
                        engine.setProperty('voice', voices[voice_options[selected_voice]].id)
    
                        # Create a temporary file to save the audio
                        temp_audio_file = "temp_audiobook.mp3"
    
                        # Split text into chunks and process
                        chunks = split_text(text)
                        total_chunks = len(chunks)
    
                        progress_bar = st.progress(0)
                        status_text = st.empty()
    
                        # Process each chunk
                        for i, chunk in enumerate(chunks, 1):
                            status_text.text(f"Processing chunk {i}/{total_chunks}...")
                            engine.save_to_file(chunk, temp_audio_file)
                            engine.runAndWait()
                            progress_bar.progress(i / total_chunks)
    
                        # Read the generated audio file
                        with open(temp_audio_file, "rb") as audio_file:
                            audio_bytes = audio_file.read()
    
                        st.success("Audiobook created successfully!")
                        st.audio(audio_bytes, format='audio/mp3')
    
                        # Provide download button
                        st.download_button(
                            label="Download Audiobook",
                            data=audio_bytes,
                            file_name="audiobook.mp3",
                            mime="audio/mpeg"
                        )
    
                        # Clean up temporary file
                        os.remove(temp_audio_file)
    
                except Exception as e:
                    st.error(f"An error occurred: {str(e)}")
            else:
                st.warning("Please provide text before converting.")

    # New feature: Book Outline Generator
    elif choice == "Book Outline Generator":
        st.header("Detailed Book Outline Generator")
        st.write("Create a comprehensive, chapter-by-chapter outline for your next bestseller.")

        col1, col2 = st.columns([2, 1])

        with col1:
            title = st.text_input("Book Title", placeholder="Enter your book title")
            main_idea = st.text_area("Main Idea or Concept", placeholder="Describe the central theme or concept of your book in detail", height=150)
            target_audience = st.text_input("Target Audience", placeholder="Who is your book for? Be specific about demographics, interests, etc.")
            key_themes = st.text_area("Key Themes or Topics", placeholder="List the main themes or topics you want to cover in your book", height=100)

        with col2:
            genre_type = st.radio("Genre Type:", ("Fiction", "Non-fiction"))
            if genre_type == "Fiction":
                genre = st.selectbox("Fiction Genre", fiction_genres)
                protagonist = st.text_input("Protagonist", placeholder="Describe your main character")
                setting = st.text_input("Setting", placeholder="Where and when does your story take place?")
            else:
                genre = st.selectbox("Non-fiction Genre", non_fiction_genres)
                expertise_level = st.select_slider("Reader's Expertise Level", options=["Beginner", "Intermediate", "Advanced", "Expert"])

            desired_length = st.number_input("Estimated Word Count", min_value=10000, value=80000, step=5000, help="Approximate length of your book")
            num_chapters = st.number_input("Number of Chapters", min_value=5, max_value=50, value=15)

        if st.button("Generate Detailed Outline", type="primary"):
            if title and main_idea and target_audience and genre and key_themes:
                with st.spinner("Crafting your comprehensive book outline..."):
                    if genre_type == "Fiction":
                        outline_prompt = f"""Generate an extremely detailed book outline for a {genre} novel titled '{title}'. 
                        Main idea: {main_idea}
                        Target audience: {target_audience}
                        Key themes: {key_themes}
                        Protagonist: {protagonist}
                        Setting: {setting}
                        Estimated length: {desired_length} words
                        Number of chapters: {num_chapters}

                        For each chapter, provide:
                        1. A compelling chapter title
                        2. A detailed synopsis (500-800 words)
                        3. Key plot points or events
                        4. Character development and interactions
                        5. Setting details and atmosphere
                        6. Themes explored in the chapter
                        7. Any foreshadowing or plot twists
                        8. Estimated word count for the chapter

                        Additionally, include:
                        - An engaging prologue idea
                        - A captivating epilogue concept
                        - Suggestions for potential subplots
                        - Ideas for symbolic elements or motifs throughout the book

                        Make the outline as comprehensive and detailed as possible, using the maximum available tokens."""
                    else:
                        outline_prompt = f"""Generate an extremely detailed book outline for a {genre} non-fiction book titled '{title}'. 
                        Main idea: {main_idea}
                        Target audience: {target_audience} (Expertise level: {expertise_level})
                        Key themes or topics: {key_themes}
                        Estimated length: {desired_length} words
                        Number of chapters: {num_chapters}

                        For each chapter, provide:
                        1. An informative chapter title
                        2. A detailed chapter summary (200-300 words)
                        3. Main concepts or arguments presented
                        4. Supporting evidence, data, or examples to include
                        5. Potential expert quotes or case studies to research
                        6. Practical applications or exercises for readers
                        7. Key takeaways from the chapter
                        8. Estimated word count for the chapter

                        Additionally, include:
                        - An attention-grabbing introduction outline
                        - A powerful conclusion and call-to-action outline
                        - Ideas for sidebars, infographics, or illustrations
                        - Suggestions for further reading or resources

                        Make the outline as comprehensive and detailed as possible, using the maximum available tokens."""

                    outline = together_client.chat.completions.create(
                        model="meta-llama/Meta-Llama-3.1-70B-Instruct-Turbo",
                        messages=[
                            {"role": "system", "content": "You are a professional book outliner and developmental editor with extensive experience in creating detailed, chapter-by-chapter outlines for bestselling books across various genres."},
                            {"role": "user", "content": outline_prompt}
                        ],
                        max_tokens=15000,
                        temperature=0.7
                    )

                    st.session_state.book_outline = outline.choices[0].message.content

                st.success("Comprehensive book outline generated successfully!")

                # Display the outline in a structured and visually appealing way
                st.subheader("Your Detailed Book Outline")
                outline_lines = st.session_state.book_outline.split('\n')
                chapter_count = 0
                for line in outline_lines:
                    if line.strip().startswith('Chapter'):
                        chapter_count += 1
                        st.markdown(f"<h3 style='color: #1e90ff;'>{line.strip()}</h3>", unsafe_allow_html=True)
                    elif any(section in line for section in ['Synopsis:', 'Summary:', 'Plot Points:', 'Main Concepts:', 'Character Development:', 'Supporting Evidence:', 'Setting:', 'Practical Applications:', 'Themes:', 'Key Takeaways:', 'Foreshadowing:', 'Estimated Word Count:']):
                        st.markdown(f"<h4 style='color: #32cd32;'>{line.strip()}</h4>", unsafe_allow_html=True)
                    else:
                        st.write(line.strip())

                st.info(f"Total Chapters: {chapter_count}")

                st.download_button(
                    label="Download Detailed Outline",
                    data=st.session_state.book_outline,
                    file_name=f"{title.replace(' ', '_').lower()}_detailed_outline.txt",
                    mime="text/plain",
                    key="download_detailed_outline"
                )
            else:
                st.warning("Please fill in all the required fields to generate a detailed outline.")

    # New feature: Character Development Workshop
    elif choice == "Character Development Workshop":
        st.header("Character Development Workshop")
        st.write("Bring your characters to life with our in-depth development tools.")

        col1, col2 = st.columns([1, 1])

        with col1:
            character_name = st.text_input("Character Name", placeholder="Enter character's name")
            character_role = st.selectbox("Character Role", ["Protagonist", "Antagonist", "Supporting Character", "Mentor", "Love Interest", "Sidekick"])
            character_age = st.number_input("Age", min_value=0, max_value=250, value=30)
            character_occupation = st.text_input("Occupation", placeholder="Character's job or main activity")

        with col2:
            character_background = st.text_area("Background", placeholder="Brief history or backstory", height=100)
            character_goals = st.text_area("Goals", placeholder="What does the character want to achieve?", height=80)
            character_fears = st.text_area("Fears or Weaknesses", placeholder="What holds the character back?", height=80)

        col3, col4 = st.columns([1, 1])

        with col3:
            physical_attributes = st.text_area("Physical Attributes", placeholder="Describe appearance, mannerisms, etc.", height=100)

        with col4:
            personality_traits = st.text_area("Personality Traits", placeholder="List key personality characteristics", height=100)

        if st.button("Develop Character", type="primary"):
            if character_name and character_role and character_background:
                with st.spinner("Crafting your character profile..."):
                    character_prompt = f"""
                    Create a detailed character profile for {character_name}:
                    - Age: {character_age}
                    - Occupation: {character_occupation}
                    - Role: {character_role}
                    - Background: {character_background}
                    - Goals: {character_goals}
                    - Fears/Weaknesses: {character_fears}
                    - Physical Attributes: {physical_attributes}
                    - Personality Traits: {personality_traits}

                    Expand on these details to create a rich, multi-dimensional character. Include potential character arc, quirks, and how they might interact with other characters or drive the plot forward.
                    """

                    character_profile = together_client.chat.completions.create(
                        model="meta-llama/Meta-Llama-3.1-70B-Instruct-Turbo",
                        messages=[
                            {"role": "system", "content": "You are a professional character developer for novels and screenplays, skilled in creating complex, believable characters."},
                            {"role": "user", "content": character_prompt}
                        ],
                        max_tokens=6500,
                        temperature=0.7
                    )

                    st.session_state.character_profile = character_profile.choices[0].message.content

                st.success("Character profile developed successfully!")

                # Display the character profile in a more structured way
                st.subheader(f"{character_name}'s Character Profile")
                profile_lines = st.session_state.character_profile.split('\n')
                for line in profile_lines:
                    if ':' in line:
                        key, value = line.split(':', 1)
                        st.markdown(f"**{key.strip()}:** {value.strip()}")
                    else:
                        st.write(line.strip())

                st.download_button(
                    label="Download Character Profile",
                    data=st.session_state.character_profile,
                    file_name=f"{character_name.replace(' ', '_').lower()}_profile.txt",
                    mime="text/plain",
                    key="download_character_profile"
                )
            else:
                st.warning("Please fill in at least the character's name, role, and background to generate a profile.")

    # New feature: Writing Prompts Generator
    elif choice == "Writing Prompts Generator":
        st.header("Writing Prompts Generator")
        st.write("Spark your creativity with custom writing prompts.")

        col1, col2 = st.columns([1, 1])

        with col1:
            prompt_type = st.selectbox("Prompt Type", ["General", "Sci-Fi", "Fantasy", "Romance", "Mystery", "Historical", "Horror", "Thriller"])
            prompt_length = st.slider("Prompt Complexity", min_value=1, max_value=5, value=3, help="1: Simple, 5: Elaborate")

        with col2:
            specific_elements = st.multiselect("Include Specific Elements", ["Character", "Setting", "Conflict", "Theme", "Plot Twist"])
            writing_style = st.selectbox("Writing Style", ["Any", "Descriptive", "Dialogue-heavy", "Action-packed", "Introspective", "Humorous"])

        mood = st.select_slider("Mood", options=["Dark", "Neutral", "Light"], value="Neutral")

        if st.button("Generate Writing Prompt", type="primary"):
            with st.spinner("Crafting your writing prompt..."):
                prompt_request = f"""
                Generate a {prompt_type.lower()} writing prompt.
                Complexity: {prompt_length}/5
                Include these elements: {', '.join(specific_elements)}
                Writing style: {writing_style}
                Mood: {mood}

                The prompt should inspire a short story or scene that a writer can immediately start working on.
                """

                generated_prompt = together_client.chat.completions.create(
                    model="meta-llama/Meta-Llama-3.1-70B-Instruct-Turbo",
                    messages=[
                        {"role": "system", "content": "You are a creative writing prompt generator, skilled in crafting inspiring and thought-provoking prompts for writers."},
                        {"role": "user", "content": prompt_request}
                    ],
                    max_tokens=1024,
                    temperature=0.9
                )

                st.session_state.writing_prompt = generated_prompt.choices[0].message.content

            st.success("Writing prompt generated successfully!")

            # Display the writing prompt in an attractive format
            st.markdown("### Your Writing Prompt")
            st.info(st.session_state.writing_prompt)

            col3, col4 = st.columns([1, 1])
            with col3:
                if st.button("Generate Another Prompt"):
                    st.experimental_rerun()
            with col4:
                st.download_button(
                    label="Save Prompt",
                    data=st.session_state.writing_prompt,
                    file_name="writing_prompt.txt",
                    mime="text/plain",
                    key="save_writing_prompt"
                )

    elif choice == "Interactive Character Board":
        st.header("Interactive Character Board")
    
        # Initialize session state variables
        if "character_board" not in st.session_state:
            st.session_state.character_board = []
        if "custom_stats" not in st.session_state:
            st.session_state.custom_stats = ["Intelligence", "Strength", "Speed", "Durability", "Energy Projection", "Fighting Skills"]

        # Character creation form
        st.subheader("Add New Character")
        with st.form("character_form"):
            new_name = st.text_input("Character Name")
            new_description = st.text_area("Character Description", max_chars=1000)
            new_image = st.file_uploader("Character Image", type=["jpg", "png", "jpeg"])

            st.subheader("Character Stats")
            stats = {}
            for stat in st.session_state.custom_stats:
                stats[stat] = st.slider(f"{stat}", 0, 10, 5)

            background = st.text_area("Character Background")
            abilities = st.text_area("Special Abilities")
            weaknesses = st.text_area("Weaknesses")
            relationships = st.text_area("Relationships")
            additional_info = st.text_area("Additional Information")

            submit_button = st.form_submit_button("Add Character")

        if submit_button:
            if new_name and new_description:
                character_data = {
                    "name": new_name,
                    "description": new_description,
                    "stats": stats,
                    "background": background,
                    "abilities": abilities,
                    "weaknesses": weaknesses,
                    "relationships": relationships,
                    "additional_info": additional_info,
                }
        
                if new_image:
                    image = Image.open(new_image)
                    img_byte_arr = io.BytesIO()
                    image.save(img_byte_arr, format='PNG')
                    character_data["image"] = img_byte_arr.getvalue()
        
                st.session_state.character_board.append(character_data)
        
                st.success(f"Character '{new_name}' added successfully!")
            else:
                st.warning("Please provide at least a name and description for the character.")
        
        # Stat management
        with st.expander("Manage Stats"):
            st.subheader("Current Stats")
            for stat in st.session_state.custom_stats:
                col1, col2 = st.columns([3, 1])
                col1.write(stat)
                if col2.button("Remove", key=f"remove_{stat}"):
                    st.session_state.custom_stats.remove(stat)
                    st.experimental_rerun()
        
            new_stat = st.text_input("Add New Stat")
            if st.button("Add Stat"):
                if new_stat and new_stat not in st.session_state.custom_stats:
                    st.session_state.custom_stats.append(new_stat)
                    st.experimental_rerun()
                else:
                    st.warning("Stat already exists or is empty.")

        # Display character cards
        st.subheader("Character Cards")
        if st.session_state.character_board:
            for idx, character in enumerate(st.session_state.character_board):
                with st.expander(character["name"], expanded=True):
                    col1, col2 = st.columns([1, 1])
                    with col1:
                        if "image" in character:
                            st.image(character["image"], use_column_width=True)
                        st.write(f"**Description:** {character['description']}")
                        st.write(f"**Background:** {character['background']}")
                        st.write(f"**Special Abilities:** {character['abilities']}")
                        st.write(f"**Weaknesses:** {character['weaknesses']}")
                        st.write(f"**Relationships:** {character['relationships']}")
                        if character["additional_info"]:
                            st.write(f"**Additional Info:** {character['additional_info']}")
                    with col2:
                        st.subheader("Character Stats")
                        fig = create_radar_chart(character["stats"])
                        st.plotly_chart(fig, use_container_width=True)

                    if st.button(f"Delete {character['name']}", key=f"delete_{idx}"):
                        st.session_state.character_board.pop(idx)
                        st.experimental_rerun()
        else:
            st.info("No characters added yet. Use the form above to add characters to the board.")

    elif choice == "World-Building Assistant":
        world_building_assistant()

        # Optional: Add a feature to view saved prompts
        if st.checkbox("View Saved Prompts"):
            if "saved_prompts" not in st.session_state:
                st.session_state.saved_prompts = []

            if st.session_state.writing_prompt and st.button("Add Current Prompt to Saved"):
                st.session_state.saved_prompts.append(st.session_state.writing_prompt)
                st.success("Prompt added to saved list!")

            if st.session_state.saved_prompts:
                for i, prompt in enumerate(st.session_state.saved_prompts):
                    st.text_area(f"Saved Prompt {i+1}", prompt, height=100, key=f"saved_prompt_{i}")
            else:
                st.info("No saved prompts yet. Generate and save some prompts to see them here!")

    elif choice == "Advanced Book Research":
            load_research_data()
            book_research()
    
    elif choice == "AI Script Generator":
        script_generator()
        
    elif choice == "Content Generator":
        content_generator()


    # Sidebar for quick access to generated content
    with st.sidebar:
        st.header("Quick Access")
        if "generated_book" in st.session_state:
            if st.button("View Generated Book"):
                st.session_state.current_view = "generated_book"
        if "book_outline" in st.session_state:
            if st.button("View Book Outline"):
                st.session_state.current_view = "book_outline"
        if "character_profile" in st.session_state:
            if st.button("View Character Profile"):
                st.session_state.current_view = "character_profile"
        if "writing_prompt" in st.session_state:
            if st.button("View Writing Prompt"):
                st.session_state.current_view = "writing_prompt"
        if "story_branches" in st.session_state:
            if st.button("View Interactive Story"):
                st.session_state.current_view = "interactive_story"
        if "world_elements" in st.session_state:
            if st.button("View World-Building"):
                st.session_state.current_view = "world_building"

    # Display the selected content in the main area
    if "current_view" in st.session_state:
        if st.session_state.current_view == "generated_book":
            st.header("Generated Book")
            st.text_area("Book Content", st.session_state.generated_book, height=400)
        elif st.session_state.current_view == "book_outline":
            st.header("Book Outline")
            st.text_area("Outline", st.session_state.book_outline, height=400)
        elif st.session_state.current_view == "character_profile":
            st.header("Character Profile")
            st.text_area("Profile", st.session_state.character_profile, height=400)
        elif st.session_state.current_view == "writing_prompt":
            st.header("Writing Prompt")
            st.text_area("Prompt", st.session_state.writing_prompt, height=200)
        elif st.session_state.current_view == "world_building":
            world_building_assistant()
if __name__ == "__main__":
    main()
